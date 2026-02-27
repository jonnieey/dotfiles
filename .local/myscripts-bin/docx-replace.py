import sys
from docx import Document

def replace_text_with_line_break(docx_path, old_text, new_text):
    # new_text may contain '\n' to indicate a line break
    parts = new_text.split('\\n')
    doc = Document(docx_path)
    found = False

    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            # Clear the paragraph and rebuild with line break
            paragraph.clear()
            # Add first part
            run = paragraph.add_run(parts[0])
            # Copy original formatting? python-docx doesn't preserve old formatting easily.
            # For simplicity, we assume default formatting.
            for part in parts[1:]:
                run.add_break()          # inserts <w:br/>
                run = paragraph.add_run(part)
            found = True
            # Save immediately (we'll overwrite the original after processing all occurrences)
    if found:
        doc.save(docx_path)
        return True
    else:
        return False

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python replace_in_docx.py <file.docx> <old> <new>")
        sys.exit(1)
    file, old, new = sys.argv[1], sys.argv[2], sys.argv[3]
    if not replace_text_with_line_break(file, old, new):
        print(f"Warning: '{old}' not found in {file}")
